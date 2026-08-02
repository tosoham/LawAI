"""
Documents API Endpoints

Endpoints for document drafting and analysis.
"""

import io
import logging
import re

from fastapi import APIRouter, HTTPException, status
from fastapi.responses import StreamingResponse

from models.requests import AnalyzeDocumentRequest, DraftDocumentRequest, ExportDocxRequest
from models.responses import (
    AnalyzeDocumentResponse,
    DocumentTemplatesResponse,
    DraftDocumentResponse,
)
from tools.registry import get_tool_registry

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["documents"])

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
# Headings in generated drafts are short, fully-capitalised lines.
_HEADING = re.compile(r"^[A-Z0-9][A-Z0-9 ,.\-/()\[\]']{2,80}$")


@router.post("/draft", response_model=DraftDocumentResponse)
async def draft_document(request: DraftDocumentRequest):
    """
    Draft a legal document.

    Generates professional legal documents like bail applications,
    petitions, and notices based on provided case details.

    Args:
        request: Document drafting request

    Returns:
        Drafted document with metadata

    Raises:
        HTTPException: If drafting fails
    """
    try:
        logger.info(f"Draft document request: type={request.document_type}")

        # Get tool registry
        registry = get_tool_registry()
        draft_tool = registry.get_tool("draft_document")

        if not draft_tool:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Draft document tool not available"
            )

        # Execute tool
        result = await draft_tool.safe_execute(
            document_type=request.document_type,
            case_details=request.case_details
        )

        if not result.success:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=result.error
            )

        # Format response
        return DraftDocumentResponse(
            success=True,
            document=result.data.get("document"),
            document_type=result.data.get("document_type"),
            case_details=result.data.get("case_details"),
            disclaimer=result.data.get("disclaimer")
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error drafting document: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to draft document: {e!s}"
        ) from e


@router.post("/analyze", response_model=AnalyzeDocumentResponse)
async def analyze_document(request: AnalyzeDocumentRequest):
    """
    Analyze a legal document.

    Provides comprehensive analysis including summaries, risk assessment,
    and key clause extraction.

    Args:
        request: Document analysis request

    Returns:
        Document analysis with insights

    Raises:
        HTTPException: If analysis fails
    """
    try:
        logger.info(f"Analyze document request: type={request.analysis_type}, doc_type={request.document_type}")

        # Get tool registry
        registry = get_tool_registry()
        analyze_tool = registry.get_tool("analyze_document")

        if not analyze_tool:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Analyze document tool not available"
            )

        # Execute tool
        result = await analyze_tool.safe_execute(
            document_text=request.document_text,
            analysis_type=request.analysis_type,
            document_type=request.document_type
        )

        if not result.success:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=result.error
            )

        # Format response
        return AnalyzeDocumentResponse(
            success=True,
            analysis=result.data.get("analysis"),
            analysis_type=result.data.get("analysis_type"),
            document_type=result.data.get("document_type"),
            document_length=result.data.get("document_length"),
            disclaimer=result.data.get("disclaimer")
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing document: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to analyze document: {e!s}"
        ) from e


@router.post(
    "/export/docx",
    responses={200: {"content": {DOCX_MEDIA_TYPE: {}}, "description": "Word document"}},
    summary="Export document text as .docx",
)
async def export_docx(request: ExportDocxRequest):
    """
    Render document text as a downloadable Word file.

    Takes the output of /documents/draft (or any text) and returns a .docx
    stream, so a generated draft can be filed rather than copy-pasted.
    """
    try:
        from docx import Document
        from docx.shared import Pt

        document = Document()

        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        if request.title:
            document.add_heading(request.title, level=1)

        # Blank lines separate paragraphs; keep all-caps lines as headings so
        # court formatting ("MOST RESPECTFULLY SHOWETH:") survives the round trip.
        for block in re.split(r"\n\s*\n", request.content.strip()):
            text = block.strip()
            if not text:
                continue
            if "\n" not in text and _HEADING.match(text):
                document.add_heading(text.title(), level=2)
            else:
                document.add_paragraph(text)

        if request.include_disclaimer:
            document.add_paragraph()
            note = document.add_paragraph(
                "DISCLAIMER: This document was generated by an AI system for "
                "informational purposes. Have it reviewed by a qualified legal "
                "professional before filing or relying on it."
            )
            note.runs[0].italic = True

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        filename = re.sub(r"[^A-Za-z0-9._-]+", "_", request.filename).strip("_") or "document"
        if not filename.endswith(".docx"):
            filename += ".docx"

        logger.info(f"Exported {filename} ({buffer.getbuffer().nbytes} bytes)")

        return StreamingResponse(
            buffer,
            media_type=DOCX_MEDIA_TYPE,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except Exception as e:
        logger.error(f"Error exporting document: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to export document: {e!s}",
        ) from e


@router.get("/templates", response_model=DocumentTemplatesResponse)
async def get_document_templates():
    """
    Get available document templates.

    Returns list of supported document types with descriptions.

    Returns:
        List of document templates
    """
    templates = [
        {
            "type": "bail_application",
            "name": "Bail Application",
            "description": "Application for regular or anticipatory bail under BNSS",
            "required_fields": [
                "accused_name",
                "fir_number",
                "sections",
                "facts",
                "police_station"
            ]
        },
        {
            "type": "petition",
            "name": "Petition",
            "description": "General petition for various legal matters",
            "required_fields": [
                "petitioner_name",
                "respondent_name",
                "facts",
                "relief_sought"
            ]
        },
        {
            "type": "notice",
            "name": "Legal Notice",
            "description": "Legal notice for various purposes",
            "required_fields": [
                "client_name",
                "recipient_name",
                "recipient_address",
                "subject",
                "facts",
                "demands"
            ]
        },
        {
            "type": "agreement",
            "name": "Agreement / Contract",
            "description": "Commercial agreement governed by Indian law",
            "required_fields": [
                "first_party_name",
                "second_party_name",
                "scope_of_agreement",
                "payment_terms",
                "term_details",
                "jurisdiction"
            ]
        }
    ]

    return DocumentTemplatesResponse(templates=templates)


@router.get("/health")
async def health_check():
    """
    Health check endpoint for documents service.

    Returns:
        Service status
    """
    registry = get_tool_registry()

    return {
        "status": "healthy",
        "tools_available": {
            "draft_document": registry.has_tool("draft_document"),
            "analyze_document": registry.has_tool("analyze_document")
        }
    }
